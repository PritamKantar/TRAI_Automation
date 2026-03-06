from flask import Flask, render_template, request, send_file, redirect, url_for, session
import os
import shutil
import pandas as pd
import win32com.client as win32
from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
import pythoncom

app = Flask(__name__)
app.secret_key = "trai_secret"

UPLOAD_FOLDER = "uploads"

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


# ---------------- LOGIN ----------------

USERNAME = "admin"
PASSWORD = "admin123"

@app.route("/")
def login():
    return render_template("login.html")


@app.route("/login", methods=["POST"])
def login_post():

    username = request.form["username"]
    password = request.form["password"]

    if username == USERNAME and password == PASSWORD:
        session["user"] = username
        return redirect("/automation")

    return "Invalid Login"


@app.route("/automation")
def index():

    if "user" not in session:
        return redirect("/")

    return render_template("index.html")


# ---------------- YOUR ORIGINAL CODE ----------------

@app.route("/process", methods=["POST"])
def process():

    word = request.files["word_file"]
    excel = request.files["excel_file"]
    mapping = request.files["mapping_file"]

    word_path = os.path.join(UPLOAD_FOLDER, word.filename)
    excel_path = os.path.join(UPLOAD_FOLDER, excel.filename)
    mapping_path = os.path.join(UPLOAD_FOLDER, mapping.filename)

    word.save(word_path)
    excel.save(excel_path)
    mapping.save(mapping_path)

    output_file = os.path.join(UPLOAD_FOLDER,"UPDATED_REPORT.docx")

    run_automation(word_path, excel_path, mapping_path, output_file)

    return send_file(output_file, as_attachment=True)


def run_automation(word_file, excel_file, mapping_file, output_file):
    pythoncom.CoInitialize()
    shutil.copy(word_file, output_file)

    wb = load_workbook(excel_file, data_only=True)

    mapping_df = pd.read_excel(mapping_file)
    mapping_df.columns = mapping_df.columns.str.strip()

    chart_mapping = mapping_df[mapping_df["ObjectType"] == "Chart"]

    def get_excel_value_chart(ref):

        try:
            sheet, cell = ref.split("!")
            value = wb[sheet][cell].value

            if isinstance(value, (int, float)):
                value = round(value, 2)

            return value / 100

        except:
            return None

    word = win32.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(os.path.abspath(output_file))

    chart_index = 1

    def update_chart(chart, chart_index):

        chart.ChartData.Activate()

        wb_chart = chart.ChartData.Workbook
        wb_chart.Application.Visible = False

        sheet = wb_chart.Worksheets(1)

        chart_name = f"Chart {chart_index}"

        chart_rows = chart_mapping[chart_mapping["ObjectName"] == chart_name]

        for _, row in chart_rows.iterrows():

            chart_cell = row["DestinationCell"]
            excel_ref = row["SourceCell"]

            value = get_excel_value_chart(excel_ref)

            if value is not None:
                sheet.Range(chart_cell).Value = value

        chart.Refresh()

    for shape in doc.Shapes:
        if shape.HasChart:
            update_chart(shape.Chart, chart_index)
            chart_index += 1

    for shape in doc.InlineShapes:
        if shape.HasChart:
            update_chart(shape.Chart, chart_index)
            chart_index += 1

    doc.Save()  
    doc.Close()
    word.Quit()
    pythoncom.CoUninitialize()

    text_mapping = mapping_df[mapping_df["ObjectType"] == "Text"]

    def format_value(value):

        if value is None:
            return ""

        if isinstance(value, (int, float)):
            value = round(value, 2)
            return f"{value:.2f}"

        return str(value)

    def get_excel_value_text(ref):

        try:
            sheet, cell = ref.split("!")
            value = wb[sheet][cell].value
            return format_value(value)

        except:
            return ""

    def replace_placeholders(text):

        for _, row in text_mapping.iterrows():

            placeholder = str(row["ObjectName"])
            source_ref = row["SourceCell"]

            value = get_excel_value_text(source_ref)

            text = text.replace(placeholder, value)

        return text

    def process_paragraph(paragraph):

        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        new_text = replace_placeholders(full_text)

        if new_text != full_text:

            paragraph.runs[0].text = new_text

            for run in paragraph.runs[1:]:
                run.text = ""

    doc = Document(output_file)

    for para in doc.paragraphs:
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    for textbox in doc.element.iter(qn('w:txbxContent')):

        for paragraph in textbox.iter(qn('w:p')):

            texts = list(paragraph.iter(qn('w:t')))

            for t in texts:

                if not t.text:
                    continue

                new_text = replace_placeholders(t.text)

                if new_text != t.text:
                    t.text = new_text

    doc.save(output_file)

    print("Word file updated successfully")


if __name__ == "__main__":
    app.run(debug=True)